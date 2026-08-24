#!/usr/bin/env python3
"""
取证单生成器 - 后端核心逻辑（无 GUI 依赖）

从 qzdd_generator.py (Tkinter V5) 抽取，供 pywebview 桌面版使用：
- 模板 docx 生成（页眉编号 / 表格填充 / 页脚页码字段）
- 参考表（单位全称 → 发文代字简称）加载
- 表单状态持久化（JSON）
- 编号递增
"""

import base64
import json
import re
import shutil
import subprocess
import sys
import uuid
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 路径（兼容 PyInstaller 打包）──────────────────────────────────
def _base_dir():
    """资源目录：PyInstaller 打包时用 sys._MEIPASS，否则用项目根目录"""
    if getattr(sys, "frozen", False):
        return Path(sys._MEIPASS)
    return Path(__file__).parent.parent


def _work_dir():
    """工作目录：用户输出/保存文件的位置"""
    if getattr(sys, "frozen", False):
        exe = Path(sys.executable)
        # macOS .app: .../X.app/Contents/MacOS/exe → 取 .app 所在目录
        if sys.platform == "darwin" and ".app" in str(exe):
            return exe.parents[3]
        return exe.parent
    return Path(__file__).parent.parent


PROJECT_DIR = _base_dir()
WORK_DIR = _work_dir()
TEMPLATE_PATH = PROJECT_DIR / "经济责任审计取证单.docx"
REF_XLSX_PATH = PROJECT_DIR / "ref" / "公司机构排序简称发文代字表.xlsx"
OUTPUT_DIR = WORK_DIR / "output"
SAVE_PATH = WORK_DIR / ".qzdd_save.json"
ATTACH_DIR = WORK_DIR / ".qzdd_attachments"  # 附件暂存区（未生成前的中转）

# 附件大小上限（base64 桥接传输，限制 80MB）
MAX_ATTACHMENT_BYTES = 80 * 1024 * 1024

AUDIT_TYPES = ["经责", "财务收支", "专项审计", "工程审计", "绩效审计"]

REQUIRED_FIELDS = {
    "project_name": "项目名称",
    "unit_name": "被审计单位名称",
    "doc_no": "取证单编号",
    "audit_year": "审计项目计划年度",
    "unit_abbr": "被审计单位简称",
}

# ── OOXML helpers ────────────────────────────────────────────────────


def _border_xml():
    return parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f"</w:tcBorders>"
    )


def _add_field_to_paragraph(paragraph, field_type):
    """向段落添加 OOXML 字段（PAGE 或 NUMPAGES）"""
    from docx.oxml import OxmlElement

    def _make_fc(fld_type):
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), fld_type)
        return el

    r = paragraph.add_run()
    r._r.append(_make_fc("begin"))
    r = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = f" {field_type} "
    instr.set(qn("xml:space"), "preserve")
    r._r.append(instr)
    r = paragraph.add_run()
    r._r.append(_make_fc("separate"))
    r = paragraph.add_run("1")
    r.font.size = Pt(11)
    r = paragraph.add_run()
    r._r.append(_make_fc("end"))


def fix_footer_numpages(doc):
    """修复页脚：用纯 API 重建 '第X页，共Y页'，包含 PAGE 和 NUMPAGES 字段"""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = 1  # 居中

    r = p.add_run("第")
    r.font.size = Pt(11)
    _add_field_to_paragraph(p, "PAGE")
    r = p.add_run("页，共")
    r.font.size = Pt(11)
    _add_field_to_paragraph(p, "NUMPAGES")
    r = p.add_run("页")
    r.font.size = Pt(11)


def set_cell_text(cell, text: str, para_idx: int = 0):
    p = cell.paragraphs[para_idx] if para_idx < len(cell.paragraphs) else cell.add_paragraph()
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(_border_xml())


# ── 参考表 ──────────────────────────────────────────────────────────
def load_unit_ref(xlsx_path: Path = REF_XLSX_PATH):
    """返回 {全称: 发文代字}"""
    mapping = {}
    if not xlsx_path.exists():
        print(f"[警告] 参考表不存在: {xlsx_path}")
        return mapping
    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(xlsx_path), read_only=True)
        ws = wb[wb.sheetnames[0]]
        for row in ws.iter_rows(min_row=5, values_only=True):
            if row[0] is None:
                continue
            full = str(row[1]).strip() if row[1] else ""
            abbr = str(row[3]).strip() if row[3] else ""
            if full and abbr:
                mapping[full] = abbr
        wb.close()
        print(f"[参考表] 已加载 {len(mapping)} 家单位简称")
    except Exception as e:
        print(f"[警告] 参考表加载失败: {e}")
    return mapping


# ── 编号 ────────────────────────────────────────────────────────────
def make_header_code(data: dict) -> str:
    """拼接页眉编号 SJ{年度}-{类型}-{简称}-{编号}，信息不全时返回空串"""
    y = str(data.get("audit_year", "")).strip()
    t = str(data.get("audit_type", "")).strip()
    a = str(data.get("unit_abbr", "")).strip()
    n = str(data.get("doc_no", "")).strip()
    if y and t and a and n:
        return f"SJ{y}-{t}-{a}-{n}"
    return ""


def increment_doc_no(current: str) -> str:
    """编号 +1，保持位数补零；无法解析时返回 '01'"""
    current = (current or "").strip()
    if current.isdigit():
        return str(int(current) + 1).zfill(len(current))
    digits = "".join(ch for ch in current if ch.isdigit())
    if digits:
        return str(int(digits) + 1).zfill(len(digits))
    return "01"


# ── 校验 ────────────────────────────────────────────────────────────
def validate_data(data: dict):
    """返回错误信息列表（空列表 = 通过）"""
    errors = []
    for key, label in REQUIRED_FIELDS.items():
        if not str(data.get(key, "")).strip():
            errors.append(f"请填写{label}")
    for key, label, lo, hi in [("year", "编制年份", 1900, 2200),
                               ("month", "编制月份", 1, 12),
                               ("day", "编制日", 1, 31)]:
        v = str(data.get(key, "")).strip()
        if not v.isdigit() or not (lo <= int(v) <= hi):
            errors.append(f"{label}无效: {v or '(空)'}")
    y = str(data.get("audit_year", "")).strip()
    if y and (not y.isdigit() or len(y) != 4):
        errors.append(f"审计项目计划年度应为4位数字: {y}")
    return errors


# ── 文档生成 ────────────────────────────────────────────────────────
def append_attachment_catalog_to_cell(cell, attachment_names):
    """在摘要单元格内、用户内容之后追加附件目录"""
    if not attachment_names:
        return
    if cell.text.strip():
        cell.add_paragraph()  # 与手动输入内容之间空一行
    head = cell.add_paragraph()
    run = head.add_run("附件目录：")
    run.font.bold = True
    for i, name in enumerate(attachment_names, 1):
        cell.add_paragraph(f"附件{i}：{name}")


def generate_docx(data: dict, output_path: Path, attachment_names=None):
    doc = Document(str(TEMPLATE_PATH))
    table = doc.tables[0]

    header_code = make_header_code(data)

    # 页眉
    section = doc.sections[0]
    header = section.header
    for p in header.paragraphs:
        p.clear()
    h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    h_para.alignment = 2
    h_run = h_para.add_run(f"编号：{header_code}")
    h_run.font.size = Pt(10.5)
    h_run.font.name = "仿宋_GB2312"

    # 表格
    set_cell_text(table.rows[0].cells[3], data["project_name"])
    set_cell_text(table.rows[1].cells[3], data["unit_name"])
    set_cell_text(table.rows[2].cells[3], data["matter_name"])

    summary_cell = table.rows[3].cells[1]
    set_cell_text(summary_cell, data["summary"])
    for ci in range(1, 9):
        set_cell_borders(table.rows[3].cells[ci])
    # 附件目录追加在摘要单元格内、用户内容之后
    append_attachment_catalog_to_cell(summary_cell, attachment_names or [])

    auditors = "　".join(filter(None, [data.get("auditor1", ""), data.get("auditor2", "")]))
    set_cell_text(table.rows[4].cells[2], auditors)

    date_str = f"{data['year']}年{data['month']}月{data['day']}日"
    set_cell_text(table.rows[4].cells[7], date_str)

    fix_footer_numpages(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return header_code


def generate(data: dict):
    """完整生成流程：校验 → 生成文件夹（docx + 附件）→ 返回结果 dict

    data["attachments"]: [{"id": ..., "formal_name": ...}]，顺序即附件编号顺序
    输出：output/{编号}-{事项}/ 文件夹，内含 docx 和 附件N-正式名称.原扩展名
    """
    errors = validate_data(data)
    if errors:
        return {"ok": False, "errors": errors}

    header_code = make_header_code(data)
    matter = str(data.get("matter_name", "")).strip() or "未命名事项"
    safe = re.sub(r'[\\/:*?"<>|]', "_", matter)
    base = f"{header_code}-{safe}"
    folder = OUTPUT_DIR / base
    docx_path = folder / f"{base}.docx"

    # 解析附件（从暂存区取文件）
    attachments = data.get("attachments") or []
    staged = []
    for i, att in enumerate(attachments, 1):
        att_id = str(att.get("id", ""))
        formal = str(att.get("formal_name", "")).strip()
        src = ATTACH_DIR / att_id
        real = _resolve_staged(att_id)
        if not formal:
            return {"ok": False, "errors": [f"附件{i}的正式名称为空"]}
        if real is None:
            return {"ok": False, "errors": [f"附件{i}（{formal}）的文件已丢失，请重新添加"]}
        staged.append((real, formal))

    try:
        folder.mkdir(parents=True, exist_ok=True)
        # 清理旧的生成物，避免残留
        for old in folder.iterdir():
            if old.is_file():
                old.unlink()

        names = []
        for i, (src, formal) in enumerate(staged, 1):
            safe_name = re.sub(r'[\\/:*?"<>|]', "_", formal)
            dest = folder / f"附件{i}-{safe_name}{src.suffix}"
            shutil.copy2(str(src), str(dest))
            names.append(formal)

        generate_docx(data, docx_path, attachment_names=names)
    except Exception as e:
        return {"ok": False, "errors": [f"生成失败：{e}"]}

    return {"ok": True, "path": str(folder), "filename": docx_path.name,
            "code": header_code, "attachments": len(staged)}


# ── 附件暂存区 ──────────────────────────────────────────────────────
def _resolve_staged(att_id: str):
    """按 id 在暂存区找实际文件（文件名形如 {id}{原扩展名}）"""
    if not att_id or not re.fullmatch(r"[0-9a-f]{12}", att_id):
        return None
    matches = list(ATTACH_DIR.glob(f"{att_id}.*"))
    return matches[0] if matches else None


def _save_staged(att_id: str, filename: str, data_b64: str):
    """解码 base64 并写入暂存区，返回 (路径, 字节数)"""
    try:
        raw = base64.b64decode(data_b64, validate=True)
    except Exception:
        raise ValueError("文件数据传输损坏")
    if len(raw) > MAX_ATTACHMENT_BYTES:
        raise ValueError(f"文件超过 {MAX_ATTACHMENT_BYTES // 1024 // 1024}MB 上限")
    ext = Path(filename).suffix.lower()[:20]
    ATTACH_DIR.mkdir(parents=True, exist_ok=True)
    path = ATTACH_DIR / f"{att_id}{ext}"
    with open(path, "wb") as f:
        f.write(raw)
    return path, len(raw)


def add_attachment(filename: str, data_b64: str):
    """前端拖入文件 → 暂存，返回附件元数据"""
    att_id = uuid.uuid4().hex[:12]
    try:
        path, size = _save_staged(att_id, filename, data_b64)
    except ValueError as e:
        return {"ok": False, "error": str(e)}
    formal = Path(filename).stem.strip() or filename
    return {"ok": True, "attachment": {
        "id": att_id, "orig_name": filename,
        "formal_name": formal, "size": size,
    }}


def replace_attachment(att_id: str, filename: str, data_b64: str):
    """修订替换：保留 id 和正式名称，替换文件内容"""
    old = _resolve_staged(att_id)
    if old is None:
        return {"ok": False, "error": "原附件不存在，请删除后重新添加"}
    old.unlink()
    try:
        path, size = _save_staged(att_id, filename, data_b64)
    except ValueError as e:
        return {"ok": False, "error": str(e)}
    return {"ok": True, "attachment": {
        "id": att_id, "orig_name": filename, "size": size,
    }}


def remove_attachment(att_id: str):
    old = _resolve_staged(att_id)
    if old:
        old.unlink()
    return {"ok": True}


def clear_attachments(keep_ids=None):
    """清空暂存区（keep_ids 中的 id 保留）"""
    keep = set(keep_ids or [])
    if ATTACH_DIR.exists():
        for f in ATTACH_DIR.iterdir():
            if f.is_file() and f.stem not in keep:
                f.unlink()
    return {"ok": True}


# ── 状态持久化 ──────────────────────────────────────────────────────
FORM_KEYS = ["doc_no", "audit_year", "audit_type", "unit_abbr",
             "project_name", "unit_name", "matter_name", "summary",
             "auditor1", "auditor2", "year", "month", "day"]


def default_state():
    t = date.today()
    return {
        "doc_no": "", "audit_year": "", "audit_type": "经责",
        "unit_abbr": "", "project_name": "", "unit_name": "",
        "matter_name": "", "summary": "", "auditor1": "", "auditor2": "",
        "year": str(t.year), "month": str(t.month), "day": str(t.day),
    }


def save_state(data: dict):
    state = {k: str(data.get(k, "")) for k in FORM_KEYS}
    # 附件只存元数据（id/名称/顺序），文件本体在暂存区
    atts = []
    for att in (data.get("attachments") or []):
        if _resolve_staged(str(att.get("id", ""))):
            atts.append({
                "id": str(att.get("id", "")),
                "orig_name": str(att.get("orig_name", "")),
                "formal_name": str(att.get("formal_name", "")),
                "size": int(att.get("size", 0)),
            })
    state["attachments"] = atts
    with open(SAVE_PATH, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)
    return state


def load_state():
    if not SAVE_PATH.exists():
        state = default_state()
        state["attachments"] = []
        return state
    try:
        with open(SAVE_PATH, "r", encoding="utf-8") as f:
            state = json.load(f)
        merged = default_state()
        for k in FORM_KEYS:
            if k in state and state[k] is not None:
                merged[k] = str(state[k])
        # 附件：校验暂存文件还在，丢了就剔除并提示
        atts, missing = [], []
        for att in (state.get("attachments") or []):
            if _resolve_staged(str(att.get("id", ""))):
                atts.append({
                    "id": str(att.get("id", "")),
                    "orig_name": str(att.get("orig_name", "")),
                    "formal_name": str(att.get("formal_name", "")),
                    "size": int(att.get("size", 0)),
                })
            else:
                missing.append(att.get("formal_name", "?"))
        merged["attachments"] = atts
        if missing:
            print(f"[警告] {len(missing)} 个附件文件已丢失: {missing}")
        return merged
    except Exception as e:
        print(f"[警告] 状态加载失败: {e}")
        state = default_state()
        state["attachments"] = []
        return state


# ── 打开输出目录 ────────────────────────────────────────────────────
def open_output_dir():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if sys.platform == "win32":
        import os
        os.startfile(str(OUTPUT_DIR))  # noqa
    elif sys.platform == "darwin":
        subprocess.Popen(["open", str(OUTPUT_DIR)])
    else:
        subprocess.Popen(["xdg-open", str(OUTPUT_DIR)])
    return str(OUTPUT_DIR)
