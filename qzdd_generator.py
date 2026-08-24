#!/usr/bin/env python3
"""
经济责任审计取证单 自动生成工具 V5

V5 - 新增下一张 / 持久化保存 / 暂存功能 / 修复简称自动补全
V4 - 单位名称自动补全 / 简称移至编号信息 / 摘要增高 /
      审计人员+编号信息横向排列 / 页脚 NUMPAGES 字段修复
V3 - 被审计单位简称xlsx自动匹配 / 页眉预览只读可复制 / 日期拆年月日 /
      GUI布局修复 / 文件名=页眉编号+审计事项
V2 - 编号+年度+类型+页眉 / 审计人员拆分 / 摘要边框 / 页脚页码
V1 - 基础GUI+docx生成
"""

import json
import sys
import tkinter as tk
from tkinter import ttk, messagebox
from datetime import date
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 路径（兼容 PyInstaller 打包）──────────────────────────────────
def _base_dir():
    """PyInstaller 打包时用 sys._MEIPASS，否则用脚本所在目录"""
    if getattr(sys, "frozen", False):
        return Path(sys._MEIPASS)
    return Path(__file__).parent

def _work_dir():
    """工作目录：用户输出/保存文件的位置"""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).parent
    return Path(__file__).parent

PROJECT_DIR = _base_dir()
WORK_DIR = _work_dir()
TEMPLATE_PATH = PROJECT_DIR / "经济责任审计取证单.docx"
REF_XLSX_PATH = PROJECT_DIR / "ref" / "公司机构排序简称发文代字表.xlsx"
OUTPUT_DIR = WORK_DIR / "output"
SAVE_PATH = WORK_DIR / ".qzdd_save.json"

# ── OOXML helpers ────────────────────────────────────────────────────
WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _border_xml():
    return parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f"</w:tcBorders>"
    )


def _add_field_to_paragraph(paragraph, field_type):
    """向段落添加一个 OOXML 字段（PAGE 或 NUMPAGES），返回添加的 run 数。
    
    字段结构:  
      <w:r><w:fldChar w:fldCharType="begin"/></w:r>
      <w:r><w:instrText> FIELD_TYPE </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r>
      <w:r><w:t>1</w:t></w:r>
      <w:r><w:fldChar w:fldCharType="end"/></w:r>
    """
    from docx.oxml import OxmlElement

    def _make_fc(fld_type):
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), fld_type)
        return el

    # begin
    r = paragraph.add_run()
    r._r.append(_make_fc("begin"))
    # instrText
    r = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = f" {field_type} "
    instr.set(qn("xml:space"), "preserve")
    r._r.append(instr)
    # separate
    r = paragraph.add_run()
    r._r.append(_make_fc("separate"))
    # value (placeholder, Word 渲染时替换)
    r = paragraph.add_run("1")
    r.font.size = Pt(11)
    # end
    r = paragraph.add_run()
    r._r.append(_make_fc("end"))


def fix_footer_numpages(doc):
    """修复页脚：用纯 API 重建 '第X页，共Y页'，包含 PAGE 和 NUMPAGES 字段"""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = 1  # 居中

    # 第
    r = p.add_run("第")
    r.font.size = Pt(11)

    # PAGE 字段
    _add_field_to_paragraph(p, "PAGE")

    # 页，共
    r = p.add_run("页，共")
    r.font.size = Pt(11)

    # NUMPAGES 字段
    _add_field_to_paragraph(p, "NUMPAGES")

    # 页
    r = p.add_run("页")
    r.font.size = Pt(11)


def set_cell_text(cell, text: str, para_idx: int = 0):
    p = cell.paragraphs[para_idx] if para_idx < len(cell.paragraphs) else cell.add_paragraph()
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(_border_xml())


# ── 参考表 ──────────────────────────────────────────────────────────
def load_unit_ref(xlsx_path: Path):
    """返回 {全称: 发文代字}"""
    mapping = {}
    if not xlsx_path.exists():
        print(f"[警告] 参考表不存在: {xlsx_path}")
        return mapping
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(xlsx_path), read_only=True)
        ws = wb[wb.sheetnames[0]]
        for row in ws.iter_rows(min_row=5, values_only=True):
            if row[0] is None:
                continue
            full = str(row[1]).strip() if row[1] else ""
            abbr = str(row[3]).strip() if row[3] else ""
            if full and abbr:
                mapping[full] = abbr
        wb.close()
        print(f"[参考表] 已加载 {len(mapping)} 家单位简称")
    except Exception as e:
        print(f"[警告] 参考表加载失败: {e}")
    return mapping


# ── 文档生成 ────────────────────────────────────────────────────────
def generate_docx(data: dict, output_path: Path):
    doc = Document(str(TEMPLATE_PATH))
    table = doc.tables[0]

    header_code = f"SJ{data['audit_year']}-{data['audit_type']}-{data['unit_abbr']}-{data['doc_no']}"

    # 页眉
    section = doc.sections[0]
    header = section.header
    for p in header.paragraphs:
        p.clear()
    h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    h_para.alignment = 2
    h_run = h_para.add_run(f"编号：{header_code}")
    h_run.font.size = Pt(10.5)
    h_run.font.name = "仿宋_GB2312"

    # 表格
    set_cell_text(table.rows[0].cells[3], data["project_name"])
    set_cell_text(table.rows[1].cells[3], data["unit_name"])
    set_cell_text(table.rows[2].cells[3], data["matter_name"])

    set_cell_text(table.rows[3].cells[1], data["summary"])
    for ci in range(1, 9):
        set_cell_borders(table.rows[3].cells[ci])

    auditors = "　".join(filter(None, [data["auditor1"], data["auditor2"]]))
    set_cell_text(table.rows[4].cells[2], auditors)

    date_str = f"{data['year']}年{data['month']}月{data['day']}日"
    set_cell_text(table.rows[4].cells[7], date_str)

    # 修复页脚 NUMPAGES
    fix_footer_numpages(doc)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return header_code


# ══════════════════════════════════════════════════════════════════════
#  GUI
# ══════════════════════════════════════════════════════════════════════

class App:
    def __init__(self, root):
        self.root = root
        root.title("审计取证单生成器 V5")
        root.geometry("920x700")
        root.minsize(800, 600)

        self.unit_ref = load_unit_ref(REF_XLSX_PATH)
        self._unit_names = sorted(self.unit_ref.keys())
        self._generated_this_session = False  # 本窗口是否已生成过（控制"新增下一张"按钮显示）

        # ── 主容器 ──
        main = ttk.Frame(root, padding="14 10 14 8")
        main.pack(fill=tk.BOTH, expand=True)

        ttk.Label(main, text="审计（调查）取证单 生成器",
                  font=("", 16, "bold")).pack(anchor="w", pady=(0, 8))

        # 可滚动画布
        canvas = tk.Canvas(main, highlightthickness=0)
        scrollbar = ttk.Scrollbar(main, orient=tk.VERTICAL, command=canvas.yview)
        self.form_frame = ttk.Frame(canvas)
        self.form_frame.bind("<Configure>",
                             lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=self.form_frame, anchor="nw",
                             tags="form_window")
        canvas.bind("<Configure>",
                    lambda e: canvas.itemconfig("form_window", width=e.width))
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        def _mw(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", _mw))
        canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))
        root.bind("<Destroy>", lambda e: canvas.unbind_all("<MouseWheel>"))

        self._build_form()

        # 底部按钮栏
        btn_frame = ttk.Frame(main)
        btn_frame.pack(fill=tk.X, pady=(8, 2))

        # "新增下一张" 按钮（初始隐藏）
        self.btn_next = ttk.Button(btn_frame, text="新增下一张 ▶",
                                   command=self.new_next)
        self.btn_next.pack(side=tk.LEFT, padx=4)
        self.btn_next.pack_forget()  # 初始隐藏

        ttk.Button(btn_frame, text="生成取证单 ▶", command=self.generate).pack(
            side=tk.RIGHT, padx=4)
        ttk.Button(btn_frame, text="暂存", command=self.save_draft).pack(
            side=tk.RIGHT, padx=4)
        ttk.Button(btn_frame, text="清空", command=self.clear).pack(
            side=tk.RIGHT, padx=4)

        self.status_var = tk.StringVar(value="就绪")
        ttk.Label(main, textvariable=self.status_var,
                  relief=tk.SUNKEN, anchor=tk.W, padding=(8, 2)).pack(
            fill=tk.X, side=tk.BOTTOM, pady=(4, 0))

        # ── 加载持久化状态 ──
        self._load_state()

        # ── 窗口关闭时自动保存 ──
        self._destroyed = False
        root.protocol("WM_DELETE_WINDOW", self._on_close)

    # ── 构建表单 ──────────────────────────────────────────────────
    def _build_form(self):
        f = self.form_frame
        f.columnconfigure(0, weight=1)
        f.columnconfigure(1, weight=1)

        # ═══ 编号信息（横向排列）═══
        sec1 = ttk.LabelFrame(f, text="编号信息", padding="8 6 8 6")
        sec1.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0, 6))

        # Row 0: 取证单编号 | 审计项目类型
        r = 0
        ttk.Label(sec1, text="取证单编号：").grid(row=r, column=0, sticky="e", padx=(0, 4), pady=2)
        self.doc_no = ttk.Entry(sec1, width=16)
        self.doc_no.grid(row=r, column=1, sticky="w", pady=2)

        ttk.Label(sec1, text="审计项目类型：").grid(row=r, column=2, sticky="e", padx=(12, 4), pady=2)
        self.audit_type = ttk.Combobox(sec1, width=14,
                                       values=["经责", "财务收支", "专项审计", "工程审计", "绩效审计"])
        self.audit_type.grid(row=r, column=3, sticky="w", pady=2)
        self.audit_type.set("经责")
        r += 1

        # Row 1: 审计项目计划年度 | 被审计单位简称
        ttk.Label(sec1, text="审计项目计划年度：").grid(row=r, column=0, sticky="e", padx=(0, 4), pady=2)
        yf = ttk.Frame(sec1)
        yf.grid(row=r, column=1, sticky="w", pady=2)
        self.audit_year = ttk.Entry(yf, width=8)
        self.audit_year.pack(side=tk.LEFT)
        ttk.Label(yf, text="年", foreground="gray").pack(side=tk.LEFT, padx=2)

        ttk.Label(sec1, text="被审计单位简称：").grid(row=r, column=2, sticky="e", padx=(12, 4), pady=2)
        self.unit_abbr = ttk.Entry(sec1, width=18)
        self.unit_abbr.grid(row=r, column=3, sticky="w", pady=2)
        r += 1

        # 页眉预览（跨整行）
        ttk.Label(sec1, text="页眉编号预览：").grid(row=r, column=0, sticky="e", padx=(0, 4), pady=2)
        self.header_preview = tk.Entry(sec1, state="readonly", readonlybackground="white",
                                       font=("", 11))
        self.header_preview.grid(row=r, column=1, columnspan=3, sticky="ew", pady=2, padx=(0, 4))

        sec1.columnconfigure(3, weight=1)

        # ═══ 项目信息 ═══
        sec2 = ttk.LabelFrame(f, text="项目信息", padding="8 6 8 6")
        sec2.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(0, 6))
        sec2.columnconfigure(1, weight=1)
        self.sec2 = sec2  # 保存引用，用于自动补全弹窗定位

        r = 0
        # 被审计单位（手动输入 + 自动补全）
        ttk.Label(sec2, text="被审计（调查）单位：").grid(
            row=r, column=0, sticky="e", padx=(0, 4), pady=3)
        self.unit_name = ttk.Entry(sec2, width=50)
        self.unit_name.grid(row=r, column=1, sticky="ew", pady=3, padx=(0, 4))
        self.unit_name.bind("<KeyRelease>", self._on_unit_key)
        r += 1

        # 自动补全弹窗（内联 Listbox，在 sec2 内用 place 定位）
        self._ac_listbox = tk.Listbox(
            sec2, width=50, height=8, font=("", 10),
            selectmode=tk.SINGLE, activestyle="dotbox",
            highlightthickness=0, borderwidth=1, relief=tk.SOLID)
        self._ac_listbox.bind("<ButtonRelease-1>", self._on_autocomplete_click)
        self._ac_listbox.bind("<Return>", self._on_autocomplete_key)
        self._ac_listbox.bind("<Escape>", lambda e: self._hide_autocomplete())

        # 项目名称
        ttk.Label(sec2, text="项目名称：").grid(row=r, column=0, sticky="e", padx=(0, 4), pady=3)
        self.project_name = ttk.Entry(sec2, width=50)
        self.project_name.grid(row=r, column=1, sticky="ew", pady=3, padx=(0, 4))
        r += 1

        # 审计事项名称
        ttk.Label(sec2, text="审计（调查）事项：").grid(row=r, column=0, sticky="e", padx=(0, 4), pady=3)
        self.matter_name = ttk.Entry(sec2, width=50)
        self.matter_name.grid(row=r, column=1, sticky="ew", pady=3, padx=(0, 4))
        r += 1

        # 摘要（更高）
        ttk.Label(sec2, text="审计（调查）事项摘要：").grid(
            row=r, column=0, sticky="ne", padx=(0, 4), pady=3)
        sf = ttk.Frame(sec2)
        sf.grid(row=r, column=1, sticky="nsew", pady=3, padx=(0, 4))
        sf.columnconfigure(0, weight=1)
        sf.rowconfigure(0, weight=1)
        self.summary = tk.Text(sf, width=40, height=10, wrap=tk.WORD, font=("", 11))
        self.summary.grid(row=0, column=0, sticky="nsew")
        ss = ttk.Scrollbar(sf, orient=tk.VERTICAL, command=self.summary.yview)
        ss.grid(row=0, column=1, sticky="ns")
        self.summary.configure(yscrollcommand=ss.set)
        sec2.rowconfigure(r, weight=1)

        # 清空事项按钮
        btn_row = ttk.Frame(sec2)
        btn_row.grid(row=r + 1, column=1, sticky="e", padx=(0, 4), pady=(4, 0))
        ttk.Button(btn_row, text="清空审计事项和摘要", command=self.clear_matters).pack()

        # ═══ 审计人员（横向）═══
        sec3 = ttk.LabelFrame(f, text="审计人员", padding="8 6 8 6")
        sec3.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(0, 6))
        sec3.columnconfigure(1, weight=1)
        sec3.columnconfigure(3, weight=1)

        ttk.Label(sec3, text="审计人员1：").grid(row=0, column=0, sticky="e", padx=(0, 4), pady=3)
        self.auditor1 = ttk.Entry(sec3, width=22)
        self.auditor1.grid(row=0, column=1, sticky="w", pady=3)

        ttk.Label(sec3, text="审计人员2：").grid(row=0, column=2, sticky="e", padx=(8, 4), pady=3)
        self.auditor2 = ttk.Entry(sec3, width=22)
        self.auditor2.grid(row=0, column=3, sticky="w", pady=3)

        # ═══ 编制日期 ═══
        sec4 = ttk.LabelFrame(f, text="编制日期", padding="8 6 8 6")
        sec4.grid(row=3, column=0, columnspan=2, sticky="ew", pady=(0, 6))

        today = date.today()
        df = ttk.Frame(sec4)
        df.pack(fill=tk.X)
        ttk.Label(df, text="年：").pack(side=tk.LEFT)
        self.year_var = tk.StringVar(value=str(today.year))
        ttk.Entry(df, textvariable=self.year_var, width=6).pack(side=tk.LEFT)
        ttk.Label(df, text=" 月：").pack(side=tk.LEFT, padx=(6, 0))
        self.month_var = tk.StringVar(value=str(today.month))
        ttk.Entry(df, textvariable=self.month_var, width=4).pack(side=tk.LEFT)
        ttk.Label(df, text=" 日：").pack(side=tk.LEFT, padx=(6, 0))
        self.day_var = tk.StringVar(value=str(today.day))
        ttk.Entry(df, textvariable=self.day_var, width=4).pack(side=tk.LEFT)
        ttk.Button(df, text="今天", width=5,
                   command=lambda: self._set_today()).pack(side=tk.LEFT, padx=10)

        # ── 绑定预览更新 ──
        for w in [self.doc_no, self.audit_year, self.unit_abbr,
                  self.project_name, self.unit_name, self.matter_name]:
            w.bind("<KeyRelease>", lambda e: self._update_header_preview())
        self.audit_type.bind("<<ComboboxSelected>>",
                             lambda e: self._update_header_preview())

    # ══════════════════════════════════════════════════════════════════
    #  自动补全（内联 Listbox — macOS 兼容，不用 Toplevel/overrideredirect）
    # ══════════════════════════════════════════════════════════════════

    def _on_unit_key(self, event=None):
        """被审计单位键盘事件：自动补全 + 简称填充"""
        text = self.unit_name.get().strip()

        # 精准匹配 → 自动填简称
        if text in self.unit_ref:
            abbr = self.unit_ref[text]
            if self.unit_abbr.get().strip() != abbr:
                self.unit_abbr.delete(0, tk.END)
                self.unit_abbr.insert(0, abbr)

        # 显示/更新 自动补全弹窗
        if len(text) < 1:
            self._hide_autocomplete()
            self._update_header_preview()
            return

        matches = [(k, v) for k, v in self.unit_ref.items()
                   if text in k][:8]
        if matches:
            self._show_autocomplete(matches)
        else:
            self._hide_autocomplete()

        self._update_header_preview()

    def _show_autocomplete(self, matches):
        """内联显示自动补全列表（place 在 unit_name 下方）"""
        lb = self._ac_listbox
        lb.delete(0, tk.END)
        for full, abbr in matches:
            lb.insert(tk.END, f"{full}  ({abbr})")
        lb._ac_matches = matches
        lb.configure(height=min(len(matches), 8))

        # 定位在 unit_name 输入框正下方
        ex = self.unit_name.winfo_x()
        ey = self.unit_name.winfo_y() + self.unit_name.winfo_height() + 1
        ew = self.unit_name.winfo_width()
        lb.place(x=ex, y=ey, width=ew)
        lb.lift()

    def _hide_autocomplete(self, *_):
        self._ac_listbox.place_forget()

    def _apply_autocomplete_choice(self, full, abbr):
        """统一处理自动补全选中逻辑"""
        self.unit_name.delete(0, tk.END)
        self.unit_name.insert(0, full)
        self.unit_abbr.delete(0, tk.END)
        self.unit_abbr.insert(0, abbr)
        self._hide_autocomplete()
        self._update_header_preview()

    def _on_autocomplete_click(self, event):
        """鼠标点击弹窗列表项"""
        lb = event.widget
        if not hasattr(lb, '_ac_matches'):
            return
        idx = lb.nearest(event.y)
        if 0 <= idx < len(lb._ac_matches):
            full, abbr = lb._ac_matches[idx]
            self._apply_autocomplete_choice(full, abbr)

    def _on_autocomplete_key(self, event):
        """键盘回车选中列表当前项"""
        lb = event.widget
        sel = lb.curselection()
        if sel and hasattr(lb, '_ac_matches'):
            full, abbr = lb._ac_matches[sel[0]]
            self._apply_autocomplete_choice(full, abbr)

    # ══════════════════════════════════════════════════════════════════

    def _set_today(self):
        t = date.today()
        self.year_var.set(str(t.year))
        self.month_var.set(str(t.month))
        self.day_var.set(str(t.day))

    def _set_readonly(self, entry, text):
        entry.configure(state="normal")
        entry.delete(0, tk.END)
        entry.insert(0, text)
        entry.configure(state="readonly")

    def _update_header_preview(self):
        y = self.audit_year.get().strip()
        t = self.audit_type.get().strip()
        a = self.unit_abbr.get().strip()
        n = self.doc_no.get().strip()
        if y and t and a and n:
            self._set_readonly(self.header_preview, f"编号：SJ{y}-{t}-{a}-{n}")
        elif any([y, t, a, n]):
            self._set_readonly(self.header_preview, "（部分信息未填）")
        else:
            self._set_readonly(self.header_preview, "（填写后自动预览）")

    def _get_header_code(self):
        y = self.audit_year.get().strip()
        t = self.audit_type.get().strip()
        a = self.unit_abbr.get().strip()
        n = self.doc_no.get().strip()
        if y and t and a and n:
            return f"SJ{y}-{t}-{a}-{n}"
        return ""

    # ── 收集表单数据 ──────────────────────────────────────────────
    def _collect_data(self):
        return {
            "doc_no": self.doc_no.get().strip(),
            "audit_year": self.audit_year.get().strip(),
            "audit_type": self.audit_type.get().strip(),
            "unit_abbr": self.unit_abbr.get().strip(),
            "project_name": self.project_name.get().strip(),
            "unit_name": self.unit_name.get().strip(),
            "matter_name": self.matter_name.get().strip(),
            "summary": self.summary.get("1.0", tk.END).strip(),
            "auditor1": self.auditor1.get().strip(),
            "auditor2": self.auditor2.get().strip(),
            "year": self.year_var.get().strip(),
            "month": self.month_var.get().strip(),
            "day": self.day_var.get().strip(),
        }

    # ── 持久化 ────────────────────────────────────────────────────
    def _save_state(self):
        """保存当前表单状态到 JSON"""
        try:
            data = self._collect_data()
            with open(SAVE_PATH, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"[警告] 保存失败: {e}")

    def _load_state(self):
        """从 JSON 加载上次保存的表单状态"""
        if not SAVE_PATH.exists():
            return
        try:
            with open(SAVE_PATH, "r", encoding="utf-8") as f:
                state = json.load(f)

            self.doc_no.insert(0, state.get("doc_no", ""))
            self.audit_year.insert(0, state.get("audit_year", ""))
            if state.get("audit_type"):
                self.audit_type.set(state["audit_type"])
            self.unit_abbr.insert(0, state.get("unit_abbr", ""))
            self.project_name.insert(0, state.get("project_name", ""))
            self.unit_name.insert(0, state.get("unit_name", ""))
            self.matter_name.insert(0, state.get("matter_name", ""))
            if state.get("summary"):
                self.summary.insert("1.0", state["summary"])
            self.auditor1.insert(0, state.get("auditor1", ""))
            self.auditor2.insert(0, state.get("auditor2", ""))
            if state.get("year"):
                self.year_var.set(state["year"])
            if state.get("month"):
                self.month_var.set(state["month"])
            if state.get("day"):
                self.day_var.set(state["day"])

            self._update_header_preview()
            self.status_var.set("已恢复上次保存的内容")
            print("[持久化] 已恢复上次会话数据")
        except Exception as e:
            print(f"[警告] 加载失败: {e}")

    # ── 窗口关闭 ──────────────────────────────────────────────────
    def _on_close(self):
        self._save_state()
        self._destroyed = True
        self.root.destroy()

    # ── 暂存 ──────────────────────────────────────────────────────
    def save_draft(self):
        self._save_state()
        self.status_var.set("已暂存 ✓")

    # ── 清空事项 ──────────────────────────────────────────────────
    def clear_matters(self):
        """只清空审计事项名称和摘要，保留其他信息"""
        self.matter_name.delete(0, tk.END)
        self.summary.delete("1.0", tk.END)
        self.status_var.set("已清空审计事项和摘要")

    # ── 新增下一张 ────────────────────────────────────────────────
    def new_next(self):
        """编号+1，清空事项/摘要，其他保留 → 准备下一张取证单"""
        current = self.doc_no.get().strip()

        # 尝试递增编号（保持位数补零）
        if current.isdigit():
            n = int(current)
            n += 1
            padded = str(n).zfill(len(current))
        else:
            try:
                n = int("".join(ch for ch in current if ch.isdigit()))
                n += 1
                padded = str(n).zfill(
                    sum(1 for ch in current if ch.isdigit()))
            except (ValueError, IndexError):
                padded = "01"

        self.doc_no.delete(0, tk.END)
        self.doc_no.insert(0, padded)

        # 清空事项名称和摘要
        self.matter_name.delete(0, tk.END)
        self.summary.delete("1.0", tk.END)

        self._update_header_preview()
        self.status_var.set(f"已准备下一张: {padded}")

    # ── 生成 ──────────────────────────────────────────────────────
    def generate(self):
        data = self._collect_data()

        if not data["project_name"]:
            messagebox.showwarning("提示", "请填写项目名称"); return
        if not data["unit_name"]:
            messagebox.showwarning("提示", "请填写被审计单位名称"); return
        if not data["doc_no"]:
            messagebox.showwarning("提示", "请填写取证单编号"); return
        if not data["audit_year"]:
            messagebox.showwarning("提示", "请填写审计项目计划年度"); return
        if not data["unit_abbr"]:
            messagebox.showwarning("提示", "请填写被审计单位简称"); return

        header_code = self._get_header_code()
        matter = data["matter_name"] if data["matter_name"] else "未命名事项"
        safe = re.sub(r'[\\/:*?"<>|]', '_', matter)
        filename = f"{header_code}-{safe}.docx"
        output_path = OUTPUT_DIR / filename

        try:
            result_code = generate_docx(data, output_path)
            self.status_var.set(f"已生成: {output_path.name}")

            # 生成成功后自动暂存 + 显示"新增下一张"按钮
            self._save_state()
            if not self._generated_this_session:
                self._generated_this_session = True
                self.btn_next.pack(side=tk.LEFT, padx=4)

            messagebox.showinfo("成功",
                                f"取证单已生成：\n{output_path}\n\n"
                                f"可点击「新增下一张」继续添加")

        except Exception as e:
            messagebox.showerror("错误", f"生成失败：\n{e}")
            self.status_var.set("生成失败")

    # ── 清空全部 ──────────────────────────────────────────────────
    def clear(self):
        self.doc_no.delete(0, tk.END)
        self.audit_year.delete(0, tk.END)
        self.audit_type.set("经责")
        self.unit_name.delete(0, tk.END)
        self.unit_abbr.delete(0, tk.END)
        self.project_name.delete(0, tk.END)
        self.matter_name.delete(0, tk.END)
        self.summary.delete("1.0", tk.END)
        self.auditor1.delete(0, tk.END)
        self.auditor2.delete(0, tk.END)
        self._set_today()
        self._set_readonly(self.header_preview, "（填写后自动预览）")
        self._hide_autocomplete()
        self._generated_this_session = False
        self.btn_next.pack_forget()
        self.status_var.set("已清空")


# ══════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    root = tk.Tk()
    App(root)
    root.mainloop()
